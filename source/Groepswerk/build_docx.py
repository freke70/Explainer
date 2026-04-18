#!/usr/bin/env python3
"""Build a .docx from the groepspresentatie plan markdown."""

import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE = Path(__file__).parent
INPUT = BASE / "groepspresentatie_volledig_plan.md"
OUTPUT = BASE / "Groepspresentatie_Plan.docx"


def setup_styles(doc):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.space_before = Pt(2)

    for level in range(1, 5):
        s = doc.styles[f"Heading {level}"]
        s.font.name = "Calibri"
        s.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        if level == 1:
            s.font.size = Pt(18)
            s.paragraph_format.space_before = Pt(24)
        elif level == 2:
            s.font.size = Pt(14)
            s.paragraph_format.space_before = Pt(18)
        elif level == 3:
            s.font.size = Pt(12)
            s.paragraph_format.space_before = Pt(14)
        else:
            s.font.size = Pt(11)
            s.paragraph_format.space_before = Pt(10)


def add_styled_run(paragraph, text, bold=False, italic=False, color=None, size=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = size
    return run


def parse_inline(paragraph, text):
    pattern = re.compile(
        r'(\*\*\*(.+?)\*\*\*)'
        r'|(\*\*(.+?)\*\*)'
        r'|(\*(.+?)\*)'
        r'|(`(.+?)`)'
    )
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        if m.group(2):
            add_styled_run(paragraph, m.group(2), bold=True, italic=True)
        elif m.group(4):
            add_styled_run(paragraph, m.group(4), bold=True)
        elif m.group(6):
            add_styled_run(paragraph, m.group(6), italic=True)
        elif m.group(8):
            run = paragraph.add_run(m.group(8))
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x80, 0x20, 0x20)
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def clean_md_links(text):
    return re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)


def parse_table(lines):
    rows = []
    for line in lines:
        line = line.strip().strip('|')
        cells = [c.strip() for c in line.split('|')]
        rows.append(cells)
    rows = [r for r in rows if not all(re.match(r'^[-:]+$', c) for c in r)]
    return rows


def add_table(doc, rows):
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    rows = [r + [''] * (ncols - len(r)) for r in rows]

    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            clean = clean_md_links(cell_text)
            if i == 0:
                add_styled_run(p, clean, bold=True, size=Pt(9.5))
                shading = cell._element.get_or_add_tcPr()
                shd = shading.makeelement(qn('w:shd'), {
                    qn('w:val'): 'clear',
                    qn('w:color'): 'auto',
                    qn('w:fill'): 'E8E8F0',
                })
                shading.append(shd)
            else:
                parse_inline(p, clean)
                for run in p.runs:
                    if run.font.size is None:
                        run.font.size = Pt(9.5)

    doc.add_paragraph("")


def process_markdown(doc, md_text):
    # Strip <details> blocks entirely
    md_text = re.sub(r'<details>.*?</details>', '', md_text, flags=re.DOTALL)
    lines = md_text.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if re.match(r'^-{3,}$', stripped):
            i += 1
            continue

        # Headings
        heading_match = re.match(r'^(#{1,4})\s+(.+)', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = clean_md_links(heading_match.group(2))
            doc.add_heading(text, level=level)
            i += 1
            continue

        # Code blocks
        if stripped.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            shading = p._element.get_or_add_pPr()
            shd = shading.makeelement(qn('w:shd'), {
                qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): 'F5F5F5',
            })
            shading.append(shd)
            run = p.add_run(code_text)
            run.font.name = "Consolas"
            run.font.size = Pt(8.5)
            continue

        # Tables
        if stripped.startswith('|') and i + 1 < len(lines):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            rows = parse_table(table_lines)
            add_table(doc, rows)
            continue

        # Blockquotes
        if stripped.startswith('>'):
            quote_lines = []
            while i < len(lines) and (lines[i].strip().startswith('>') or
                                       (lines[i].strip() and quote_lines and
                                        not lines[i].strip().startswith('#') and
                                        not lines[i].strip().startswith('|') and
                                        not lines[i].strip().startswith('-'))):
                text = lines[i].strip().lstrip('>').strip()
                if text:
                    quote_lines.append(text)
                i += 1
            quote_text = ' '.join(quote_lines)
            quote_text = clean_md_links(quote_text)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.right_indent = Cm(0.5)
            pPr = p._element.get_or_add_pPr()
            borders = pPr.makeelement(qn('w:pBdr'), {})
            left = borders.makeelement(qn('w:left'), {
                qn('w:val'): 'single', qn('w:sz'): '12',
                qn('w:space'): '4', qn('w:color'): '888888',
            })
            borders.append(left)
            pPr.append(borders)
            parse_inline(p, quote_text)
            for run in p.runs:
                run.italic = True
                if run.font.color.rgb is None:
                    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            continue

        # Bullet lists
        if re.match(r'^[-*]\s', stripped):
            text = re.sub(r'^[-*]\s+', '', stripped)
            text = clean_md_links(text)
            p = doc.add_paragraph(style='List Bullet')
            parse_inline(p, text)
            i += 1
            continue

        # Numbered lists
        num_match = re.match(r'^(\d+)\.\s+', stripped)
        if num_match:
            text = stripped[num_match.end():]
            text = clean_md_links(text)
            p = doc.add_paragraph(style='List Number')
            parse_inline(p, text)
            i += 1
            continue

        # Regular paragraph
        text = clean_md_links(stripped)
        p = doc.add_paragraph()
        parse_inline(p, text)
        i += 1


def main():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    setup_styles(doc)

    # Title page
    doc.add_paragraph("")
    doc.add_paragraph("")
    title = doc.add_heading("Groepspresentatie Leergroep 2", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_styled_run(sub, "Volledig plan — aanpak, werkbladen & tijdlijn", size=Pt(14),
                   color=RGBColor(0x55, 0x55, 0x55))
    doc.add_page_break()

    # Process the markdown
    md_text = INPUT.read_text(encoding="utf-8")
    # Skip the H1 title (already on title page)
    md_text = re.sub(r'^#\s+.+\n', '', md_text, count=1)
    process_markdown(doc, md_text)

    doc.save(str(OUTPUT))
    print(f"Saved: {OUTPUT} ({OUTPUT.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
