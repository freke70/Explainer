#!/usr/bin/env python3
"""Build a single .docx from the Recht1 markdown summary files."""

import re
import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE = Path(__file__).parent
OUTPUT = BASE / "Recht1_Samenvatting.docx"

FILES = [
    "00_Leerarchitectuur.md",
    "01_Kernkader.md",
    "02_Hoven_en_Procedures.md",
    "03_Vergelijkingsmatrix.md",
    "04_Hulpverlener_in_de_Praktijk.md",
    "05_Zelftoets.md",
]


def setup_styles(doc):
    """Configure document styles for a clean academic look."""
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.space_before = Pt(2)

    for level in range(1, 5):
        name = f"Heading {level}"
        s = doc.styles[name]
        s.font.name = "Calibri"
        s.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        if level == 1:
            s.font.size = Pt(18)
            s.paragraph_format.space_before = Pt(24)
        elif level == 2:
            s.font.size = Pt(14)
            s.paragraph_format.space_before = Pt(18)
        elif level == 3:
            s.font.size = Pt(12)
            s.paragraph_format.space_before = Pt(14)
        else:
            s.font.size = Pt(11)
            s.paragraph_format.space_before = Pt(10)


def add_styled_run(paragraph, text, bold=False, italic=False, color=None, size=None):
    """Add a run with optional formatting."""
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = size
    return run


def parse_inline(paragraph, text):
    """Parse inline markdown (bold, italic, bold+italic, code) into runs."""
    pattern = re.compile(
        r'(\*\*\*(.+?)\*\*\*)'   # ***bold italic***
        r'|(\*\*(.+?)\*\*)'       # **bold**
        r'|(\*(.+?)\*)'           # *italic*
        r'|(`(.+?)`)'             # `code`
    )
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])

        if m.group(2):  # bold+italic
            add_styled_run(paragraph, m.group(2), bold=True, italic=True)
        elif m.group(4):  # bold
            add_styled_run(paragraph, m.group(4), bold=True)
        elif m.group(6):  # italic
            add_styled_run(paragraph, m.group(6), italic=True)
        elif m.group(8):  # code
            run = paragraph.add_run(m.group(8))
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x80, 0x20, 0x20)
        pos = m.end()

    if pos < len(text):
        paragraph.add_run(text[pos:])


def clean_md_links(text):
    """Strip markdown link syntax [text](url) -> text."""
    return re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)


def strip_details_blocks(text):
    """Remove <details> blocks entirely — images serve as the primary visual in docx."""
    text = re.sub(r'<details>.*?</details>', '', text, flags=re.DOTALL)
    return text


def parse_table(lines):
    """Parse markdown table lines into list of rows (list of cell strings)."""
    rows = []
    for line in lines:
        line = line.strip().strip('|')
        cells = [c.strip() for c in line.split('|')]
        rows.append(cells)

    # Remove separator row (--- lines)
    rows = [r for r in rows if not all(re.match(r'^[-:]+$', c) for c in r)]
    return rows


def add_table(doc, rows):
    """Add a formatted table to the document."""
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    rows = [r + [''] * (ncols - len(r)) for r in rows]

    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            clean = clean_md_links(cell_text)
            if i == 0:  # Header row
                add_styled_run(p, clean, bold=True, size=Pt(9.5))
                shading = cell._element.get_or_add_tcPr()
                shd = shading.makeelement(qn('w:shd'), {
                    qn('w:val'): 'clear',
                    qn('w:color'): 'auto',
                    qn('w:fill'): 'E8E8F0',
                })
                shading.append(shd)
            else:
                parse_inline(p, clean)
                for run in p.runs:
                    if run.font.size is None:
                        run.font.size = Pt(9.5)

    doc.add_paragraph("")  # spacing


def add_image(doc, img_path):
    """Add an image centered, scaled to fit."""
    if not img_path.exists():
        p = doc.add_paragraph()
        add_styled_run(p, f"[Afbeelding niet gevonden: {img_path.name}]", italic=True,
                       color=RGBColor(0x99, 0x33, 0x33))
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(img_path), width=Inches(5.5))


def process_markdown(doc, md_text, base_path):
    """Process markdown text and add content to the document."""
    md_text = strip_details_blocks(md_text)
    lines = md_text.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip navigation tables at end of files
        if stripped.startswith('| Vorig') or stripped.startswith('| Terug naar'):
            while i < len(lines) and lines[i].strip().startswith('|'):
                i += 1
            continue

        # Skip "## Navigatie" heading
        if stripped == '## Navigatie':
            i += 1
            continue

        # Empty line
        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if re.match(r'^-{3,}$', stripped):
            i += 1
            continue

        # Headings
        heading_match = re.match(r'^(#{1,4})\s+(.+)', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = clean_md_links(heading_match.group(2))
            doc.add_heading(text, level=level)
            i += 1
            continue

        # Images
        img_match = re.match(r'^!\[.*?\]\((.+?)\)', stripped)
        if img_match:
            img_rel = img_match.group(1)
            img_path = base_path / img_rel
            add_image(doc, img_path)
            i += 1
            continue

        # Code blocks (```...```) — render as formatted block
        if stripped.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```

            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            shading = p._element.get_or_add_pPr()
            shd = shading.makeelement(qn('w:shd'), {
                qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): 'F5F5F5',
            })
            shading.append(shd)
            run = p.add_run(code_text)
            run.font.name = "Consolas"
            run.font.size = Pt(8.5)
            continue

        # Tables
        if stripped.startswith('|') and i + 1 < len(lines):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            rows = parse_table(table_lines)
            add_table(doc, rows)
            continue

        # Blockquotes
        if stripped.startswith('>'):
            quote_lines = []
            while i < len(lines) and (lines[i].strip().startswith('>') or
                                       (lines[i].strip() and quote_lines and
                                        not lines[i].strip().startswith('#') and
                                        not lines[i].strip().startswith('|'))):
                text = lines[i].strip().lstrip('>').strip()
                if text:
                    quote_lines.append(text)
                i += 1

            quote_text = ' '.join(quote_lines)
            quote_text = clean_md_links(quote_text)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.right_indent = Cm(0.5)
            pPr = p._element.get_or_add_pPr()
            borders = pPr.makeelement(qn('w:pBdr'), {})
            left = borders.makeelement(qn('w:left'), {
                qn('w:val'): 'single', qn('w:sz'): '12',
                qn('w:space'): '4', qn('w:color'): '888888',
            })
            borders.append(left)
            pPr.append(borders)
            parse_inline(p, quote_text)
            for run in p.runs:
                run.italic = True
                if run.font.color.rgb is None:
                    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            continue

        # Bullet lists
        if re.match(r'^[-*]\s', stripped):
            text = re.sub(r'^[-*]\s+', '', stripped)
            text = clean_md_links(text)
            p = doc.add_paragraph(style='List Bullet')
            parse_inline(p, text)
            i += 1
            continue

        # Numbered lists
        num_match = re.match(r'^(\d+)\.\s+', stripped)
        if num_match:
            text = stripped[num_match.end():]
            text = clean_md_links(text)
            p = doc.add_paragraph(style='List Number')
            parse_inline(p, text)
            i += 1
            continue

        # Checkbox lists
        if re.match(r'^- \[[ x]\]', stripped):
            text = re.sub(r'^- \[[ x]\]\s*', '', stripped)
            text = clean_md_links(text)
            checkbox = '\u2610 ' if '[ ]' in line else '\u2611 '
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.add_run(checkbox)
            parse_inline(p, text)
            i += 1
            continue

        # Regular paragraph
        text = clean_md_links(stripped)
        p = doc.add_paragraph()
        parse_inline(p, text)
        i += 1


def main():
    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    setup_styles(doc)

    # Title page
    doc.add_paragraph("")
    doc.add_paragraph("")
    title = doc.add_heading("Staatsstructuur, Justitie & Rechtsbronnen", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_styled_run(sub, "Deel 1 — Studiesamenvatting", size=Pt(14),
                   color=RGBColor(0x55, 0x55, 0x55))
    doc.add_paragraph("")
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_styled_run(note, "Inleiding tot het Belgisch rechtssysteem voor hulpverleners", italic=True,
                   size=Pt(11), color=RGBColor(0x77, 0x77, 0x77))
    doc.add_page_break()

    # Process each file
    for idx, fname in enumerate(FILES):
        fpath = BASE / fname
        if not fpath.exists():
            print(f"  Skipping {fname} (not found)")
            continue

        print(f"  Processing {fname}...")
        md_text = fpath.read_text(encoding="utf-8")
        process_markdown(doc, md_text, BASE)

        # Small spacing between sections instead of page breaks
        if idx < len(FILES) - 1:
            doc.add_paragraph("")

    doc.save(str(OUTPUT))
    print(f"\nSaved: {OUTPUT} ({OUTPUT.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
